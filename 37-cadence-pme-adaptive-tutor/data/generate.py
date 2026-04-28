"""CADENCE — Real-Data Adaptive PME Tutor synthetic data generator.

Produces (all under data/):
  courses.json          - 3 demo courses (Logistics Principles Paper,
                          Sergeants Course (Resident), MOS 0411 Sustainment
                          Pipeline). Each course is anchored to the relevant
                          NAVMC 3500-series Training & Readiness Manual (or
                          MCO 1553.4B PME framework where applicable) with
                          T&R event codes.
  students.json         - 3 student profiles with submission history + forum
                          post excerpts + per-course pairings
  doctrine_index.json   - 30 doctrine entries keyed by citation (MCWP 4-11,
                          MCRP 3-40D, MCO 5400.18, etc.) with section abstracts
  cached_briefs.json    - 3 student-course combinations pre-analyzed
  rubrics/*.xlsx        - rubrics for each demo course (xlsx via openpyxl)
  assignments/*.docx    - assignment instructions in .docx form (synthetic)
  submissions/*.docx    - synthetic student submissions (.docx)

Cites the LOGCOM portal datasets in the module docstring (verbatim names):
    1. "LMS Course data sets" — real USMC .mbz Moodle 4.5+ exports
       (anonymized users, course logs, discussions, structure)
    2. "Student Written Assignment Examples" — PDF assignments + xlsx rubric
       + docx instructions

Both NEW on the LOGCOM AI Forum Hackathon portal as of 2026-04. Plug-in path
in data/load_real.py.

Records governance for the synthetic training records produced here:
**Privacy Act of 1974 (5 U.S.C. § 552a)** and **DoDI 1322.35 "Military
Education Records"** — NOT FERPA. FERPA governs K-12 / civilian higher
education; active-duty military training records are governed by the
Privacy Act and DoDI 1322.35.

Seeded random.Random(1776) for full reproducibility.
NEVER use real Marine training data — all names, EDIPIs, posts are synthetic.
"""
from __future__ import annotations

import json
import random
from datetime import datetime, timedelta, timezone
from pathlib import Path

ROOT = Path(__file__).resolve().parent
ASSIGNMENTS_DIR = ROOT / "assignments"
SUBMISSIONS_DIR = ROOT / "submissions"
RUBRICS_DIR = ROOT / "rubrics"
for d in (ASSIGNMENTS_DIR, SUBMISSIONS_DIR, RUBRICS_DIR):
    d.mkdir(parents=True, exist_ok=True)


# ---------------------------------------------------------------------------
# Course catalog (3 demo courses)
# ---------------------------------------------------------------------------

GOVERNING_AUTHORITY = (
    "Privacy Act of 1974 (5 U.S.C. § 552a) and DoDI 1322.35 "
    "\"Military Education Records\""
)


COURSES = [
    {
        "id": "log_principles_paper",
        "name": "Logistics Principles Paper",
        "code": "LOG-PRIN-26",
        "school": "MAGTF Logistics Officers Course",
        "description": (
            "Mid-career logistics officer writing assignment. Maps directly to the "
            "Student Written Assignment Examples dataset on the LOGCOM portal "
            "(assignment .docx + rubric .xlsx + sample submissions)."
        ),
        "primary_doctrine": ["MCWP 4-11", "MCRP 4-11.1F", "MCDP 4", "MCWP 4-11.7"],
        "tr_manual": "NAVMC 3500.58 — Logistics Training & Readiness Manual",
        "tr_manual_short": "NAVMC 3500.58",
        "tr_event_codes": ["LOG-DIST-2001", "LOG-PLAN-2003", "LOG-SUS-3001"],
        "governing_authority": GOVERNING_AUTHORITY,
        "rubric_axes": [
            "doctrinal_grounding",
            "logistics_principles_application",
            "writing_clarity",
            "critical_thinking",
            "rubric_alignment",
        ],
    },
    {
        "id": "sergeants_course",
        "name": "Sergeants Course (Resident, SNCO Academy)",
        "code": "SGTS-CRS-2-26",
        "school": "Staff NCO Academy, MCB Quantico VA",
        "description": (
            "Resident NCO PME for E-5 Sergeants per MCO 1553.4B (PME framework). "
            "The legacy Squad Leader Distance Learning content has been folded "
            "into resident Sergeants Course (E-5) and Corporals Course (E-4); "
            "SQDL is no longer a current standalone PME pathway. Discussion "
            "forums + 4 written checks."
        ),
        "primary_doctrine": ["MCRP 3-10A.4", "MCDP 1", "MCWP 3-01", "MCRP 6-11D"],
        "tr_manual": (
            "MCO 1553.4B (PME Framework) and DoDI 1322.35 \"Military Education\" "
            "— anchored to NAVMC 3500.18 Infantry T&R for tactical events"
        ),
        "tr_manual_short": "MCO 1553.4B / NAVMC 3500.18",
        "tr_event_codes": ["INF-MAN-1001", "INF-OPS-2001", "INF-PAT-2002"],
        "governing_authority": GOVERNING_AUTHORITY,
        "rubric_axes": [
            "doctrinal_grounding",
            "leadership_judgment",
            "writing_clarity",
            "tactical_decision_making",
            "rubric_alignment",
        ],
    },
    {
        "id": "mos_0411_sustainment",
        "name": "MOS 0411 Sustainment Pipeline",
        "code": "MOS-0411-26",
        "school": "Marine Corps Logistics Operations Group",
        "description": (
            "Maintenance Management Specialist (0411) PMOS pipeline — class V/IX "
            "planning, preventive-maintenance scheduling, GCSS-MC interaction."
        ),
        "primary_doctrine": ["MCO 4790.25", "MCWP 4-11.4", "MCO 5400.18", "MCRP 3-40D"],
        "tr_manual": "NAVMC 3500.58 — Logistics Training & Readiness Manual",
        "tr_manual_short": "NAVMC 3500.58",
        "tr_event_codes": ["LOG-MAINT-1001", "LOG-MAINT-2003", "LOG-DIST-2001"],
        "governing_authority": GOVERNING_AUTHORITY,
        "rubric_axes": [
            "doctrinal_grounding",
            "maintenance_planning_factors",
            "writing_clarity",
            "systems_thinking",
            "rubric_alignment",
        ],
    },
]


# ---------------------------------------------------------------------------
# Doctrine index (30 entries) — keyed by citation
# ---------------------------------------------------------------------------

DOCTRINE_INDEX = {
    "MCWP 4-11": {
        "title": "Tactical-Level Logistics",
        "section_abstracts": {
            "Ch 1": "Logistics functions defined: supply, maintenance, transportation, general engineering, health services, services.",
            "Ch 3": "Class V (ammunition) planning factors and consumption rates by unit type and operation tempo.",
            "Ch 5": "Sustainment estimate process; tying logistic estimate to commander's COA selection.",
        },
    },
    "MCRP 4-11.1F": {
        "title": "Tactical-Level Logistics — Reference Publication",
        "section_abstracts": {
            "App A": "Standard sustainment planning factors; Class I/III/V/IX rates for MEU/MEB/MEF.",
            "App C": "Throughput and reception, staging, onward movement, integration (RSOI) calculations.",
        },
    },
    "MCDP 4": {
        "title": "Logistics",
        "section_abstracts": {
            "Ch 1": "Logistics as the bridge between national resources and tactical employment.",
            "Ch 2": "The seven principles of logistics: responsiveness, simplicity, flexibility, economy, attainability, sustainability, survivability.",
            "Ch 3": "Operational level: integration of logistics with the scheme of maneuver.",
        },
    },
    "MCWP 4-11.7": {
        "title": "MAGTF Supply Operations",
        "section_abstracts": {
            "Ch 2": "Class I subsistence planning factors and field-feeding requirements.",
            "Ch 4": "Class IX repair parts demand support and the GCSS-MC interface.",
        },
    },
    "MCRP 3-10A.4": {
        "title": "Marine Rifle Squad",
        "section_abstracts": {
            "Ch 2": "Squad organization, weapons, and roles. Fire team functions.",
            "Ch 4": "Squad in the offense — formations, control measures, fire and movement.",
            "Ch 5": "Squad in the defense — sectors of fire, EAs, withdrawal under pressure.",
        },
    },
    "MCDP 1": {
        "title": "Warfighting",
        "section_abstracts": {
            "Ch 1": "Nature of war: friction, uncertainty, fluidity, disorder, complexity.",
            "Ch 4": "Conduct of war: maneuver warfare philosophy; tempo, focus, surprise, boldness, combined arms.",
        },
    },
    "MCWP 3-01": {
        "title": "Offensive and Defensive Tactics",
        "section_abstracts": {
            "Ch 2": "Offensive operations: forms of maneuver, types of attack.",
            "Ch 5": "Defensive operations: defense in depth, mutually supporting positions.",
        },
    },
    "MCRP 6-11D": {
        "title": "Sustaining the Transformation",
        "section_abstracts": {
            "Ch 3": "Roles of the NCO in mentorship and standards enforcement.",
            "Ch 5": "Five horizontal themes — fidelity, fortitude, initiative, perseverance, integrity.",
        },
    },
    "MCO 4790.25": {
        "title": "Ground Equipment Records Procedures",
        "section_abstracts": {
            "Ch 2": "Maintenance management cycle; PMCS, scheduled maintenance, modifications.",
            "Ch 4": "Equipment status reporting; readiness reportable equipment list (RREL).",
        },
    },
    "MCWP 4-11.4": {
        "title": "Maintenance Operations",
        "section_abstracts": {
            "Ch 2": "Echelons of maintenance; field, sustainment, and depot levels.",
            "Ch 3": "Recovery and evacuation; cross-leveling parts within maintenance battalion.",
        },
    },
    "MCO 5400.18": {
        "title": "Marine Corps Logistics Command Structure",
        "section_abstracts": {
            "Ch 1": "LOGCOM mission, organization, and supporting establishment relationships.",
            "Ch 3": "Depot-level support for ground equipment; rebuild and reset programs.",
        },
    },
    "MCRP 3-40D": {
        "title": "Sustainment Operations",
        "section_abstracts": {
            "Ch 2": "Distribution architecture: hubs, nodes, lines of communication.",
            "Para 5": "Operational reach and culmination — when sustainment becomes the limiting factor.",
        },
    },
    "MCDP 5": {
        "title": "Planning",
        "section_abstracts": {
            "Ch 2": "Marine Corps Planning Process (MCPP) — six-step cyclical model.",
            "Ch 3": "Single-battle concept; integrating shaping, decisive, and sustaining actions.",
        },
    },
    "MCDP 6": {
        "title": "Command and Control",
        "section_abstracts": {
            "Ch 3": "Mission-type orders; centralized planning, decentralized execution.",
            "Ch 4": "Information management and the commander's critical information requirements (CCIR).",
        },
    },
    "MCWP 5-10": {
        "title": "Marine Corps Planning Process",
        "section_abstracts": {
            "Ch 2": "Step 1 — Problem framing.",
            "Ch 4": "Step 4 — Course of action wargaming.",
        },
    },
    "MCO P3500.72A": {
        "title": "Marine Corps Ground Training and Readiness",
        "section_abstracts": {
            "Ch 1": "T&R program structure and event-based training.",
            "Ch 3": "Combat readiness percentage (CRP) calculation.",
        },
    },
    "MCWP 3-32": {
        "title": "Marine Air-Ground Task Force Information Operations",
        "section_abstracts": {
            "Ch 2": "IO core competencies and the integration with kinetic operations.",
        },
    },
    "MCRP 3-30.4": {
        "title": "Marine Corps Operations",
        "section_abstracts": {
            "Ch 4": "Sustainment warfighting function across the range of military operations.",
        },
    },
    "MCWP 3-40.1": {
        "title": "MAGTF Command and Control",
        "section_abstracts": {
            "Ch 5": "Command relationships: OPCON, TACON, ADCON, support relationships.",
        },
    },
    "MCO 1500.59A": {
        "title": "Marine Corps Distance Learning Program",
        "section_abstracts": {
            "Ch 2": "DL course completion standards and credit toward PME requirements.",
        },
    },
    "MCO 1553.3B": {
        "title": "Unit Training Management",
        "section_abstracts": {
            "Ch 4": "Long-range training plan construction and the unit-level T&R event sequencing.",
        },
    },
    "MCRP 4-11.3K": {
        "title": "Convoy Operations Handbook",
        "section_abstracts": {
            "Ch 2": "Convoy planning factors: serial, march unit, intervals, speed.",
            "Ch 5": "Hardening and reaction drills under contact.",
        },
    },
    "MCWP 4-12": {
        "title": "Operational-Level Logistics",
        "section_abstracts": {
            "Ch 3": "Theater distribution and the joint deployment distribution operations center (JDDOC).",
        },
    },
    "MCO 4400.150": {
        "title": "Consumer-Level Supply Policy",
        "section_abstracts": {
            "Ch 3": "Stock listing, requisitioning, and on-hand balance reconciliation.",
        },
    },
    "MCO 4400.16K": {
        "title": "Uniform Materiel Movement and Issue Priority System",
        "section_abstracts": {
            "Ch 2": "UMMIPS priority designators and required delivery dates (RDDs).",
        },
    },
    "MCRP 4-11.4A": {
        "title": "Combat Service Support for Sustained Operations",
        "section_abstracts": {
            "Ch 3": "Sustainment estimate: building the days-of-supply (DOS) calculation.",
        },
    },
    "MCO 1700.29": {
        "title": "Marine Corps Lifelong Learning Program",
        "section_abstracts": {
            "Ch 2": "Voluntary education and PME credit articulation.",
        },
    },
    "MCO 5530.14A": {
        "title": "Marine Corps Physical Security Program",
        "section_abstracts": {
            "Ch 5": "Arms, ammunition, and explosives (AA&E) facility standards.",
        },
    },
    "MCO P5102.1B": {
        "title": "Marine Corps Mishap Investigation",
        "section_abstracts": {
            "Ch 2": "Class A/B/C mishap definitions and reporting timelines.",
        },
    },
    "MCO 3501.36": {
        "title": "Marine Corps Combat Readiness Reporting",
        "section_abstracts": {
            "Ch 2": "DRRS-MC reporting cycle and the C-level rollup methodology.",
        },
    },
    "MCWP 3-35.3": {
        "title": "Military Operations on Urbanized Terrain",
        "section_abstracts": {
            "Ch 5": "Building entry, room clearing, fatal funnel mitigation.",
        },
    },
}


# ---------------------------------------------------------------------------
# Synthetic students (3 profiles), each paired with one of the 3 courses
# ---------------------------------------------------------------------------

STUDENTS = [
    {
        "student_id": "M001",
        "name": "Capt Alvarez, J.",
        "rank": "Capt",
        "edipi_synth": "1234567890",
        "primary_course_id": "log_principles_paper",
        "profile": "developing",  # mid-career officer, decent doctrine, weak synthesis
        "submission_history": [
            {"course_id": "log_principles_paper", "assignment": "Logistics Principles Paper Draft 1", "grade": 78, "submitted": "2026-04-08"},
            {"course_id": "log_principles_paper", "assignment": "Discussion 2 Reflection", "grade": 82, "submitted": "2026-04-15"},
            {"course_id": "log_principles_paper", "assignment": "Sustainment Estimate Worksheet", "grade": 75, "submitted": "2026-04-20"},
        ],
    },
    {
        "student_id": "M002",
        "name": "Sgt Brennan, K.",
        "rank": "Sgt",
        "edipi_synth": "2345678901",
        "primary_course_id": "sergeants_course",
        "profile": "rising_star",  # NCO with excellent tactical thinking, weaker writing
        "submission_history": [
            {"course_id": "sergeants_course", "assignment": "Squad in the Defense — Written Check", "grade": 88, "submitted": "2026-04-10"},
            {"course_id": "sergeants_course", "assignment": "Mission-Type Orders Reflection", "grade": 91, "submitted": "2026-04-17"},
            {"course_id": "sergeants_course", "assignment": "Casualty Evac Plan", "grade": 84, "submitted": "2026-04-22"},
        ],
    },
    {
        "student_id": "M003",
        "name": "Cpl Chen, L.",
        "rank": "Cpl",
        "edipi_synth": "3456789012",
        "primary_course_id": "mos_0411_sustainment",
        "profile": "needs_remediation",  # PMOS pipeline student behind on doctrine grounding
        "submission_history": [
            {"course_id": "mos_0411_sustainment", "assignment": "PMCS Cycle Quiz", "grade": 64, "submitted": "2026-04-09"},
            {"course_id": "mos_0411_sustainment", "assignment": "GCSS-MC Practical Exercise", "grade": 58, "submitted": "2026-04-16"},
            {"course_id": "mos_0411_sustainment", "assignment": "Class IX Demand Forecast", "grade": 71, "submitted": "2026-04-21"},
        ],
    },
]


# Forum post excerpts per student (last 10 each, varied cognitive depth)
FORUM_POSTS = {
    "M001": [
        {"thread": "Apply MCDP 4 principles to a 30-day MEU sustainment", "depth": "application", "body": "Applying responsiveness, the MEU's CLB needs the LCE to push 5-day pulses rather than waiting for pull from the BLT. Per MCDP 4 Ch 2 the responsiveness principle implies forward-positioned stocks."},
        {"thread": "Class V planning factors — discussion", "depth": "recall", "body": "MCWP 4-11 Ch 3 has the consumption tables. I used the offensive operations baseline."},
        {"thread": "Where does sustainment estimate fit in MCPP?", "depth": "application", "body": "It feeds COA wargaming in step 4 of MCPP. The logistic estimate has to be live so the wargame doesn't pick a COA we cannot sustain."},
        {"thread": "Critique: simplicity vs flexibility in MCDP 4", "depth": "analysis", "body": "Simplicity argues for fewer SKU lines and standardized loadouts. Flexibility argues for tailored task-organized packages. Tension is real and the doctrine glosses it. I'd weight simplicity at the BLT level and flexibility at the LCE level."},
        {"thread": "RSOI throughput — calculation walkthrough", "depth": "application", "body": "Used MCRP 4-11.1F App C. My APOD throughput calc came out to 1,200 STON/day given two C-17 spots and a 22-hr ops day."},
        {"thread": "How does GCSS-MC handle cross-leveling?", "depth": "recall", "body": "Through the requisition module I think. We covered this in PMOS school but I have to look it up again."},
        {"thread": "Logistics rebellion — when does the tail wag the dog?", "depth": "analysis", "body": "Sustainment becomes operational reach. If the planner doesn't model culmination per MCRP 3-40D para 5, the COA will overrun its DOS and the tail dictates the new mission."},
        {"thread": "MEU sustainment estimate — peer review", "depth": "application", "body": "Peer reviewed S04's draft. Suggested rebuilding Class III(B) on JP-8 consumption per fuel-burner type, not per platform-day."},
        {"thread": "What's the worst sustainment failure you've seen written up?", "depth": "recall", "body": "The Fallujah sustainment AAR — Class IX backlog hurt us at the company-level."},
        {"thread": "MCDP 4 principle of survivability — apply to drone-saturated environment", "depth": "application", "body": "Survivability now demands distributed sustainment — fewer big nodes, more small nodes. Not in the doctrine yet."},
    ],
    "M002": [
        {"thread": "Squad in the defense — sectors of fire", "depth": "synthesis", "body": "Combining MCRP 3-10A.4 Ch 5 with MCWP 3-01 Ch 5: the squad's sector should overlap with the adjacent squad's by 50m to deny the enemy a seam, but the leader has to publish the seam owner explicitly or both squads will assume the other has it."},
        {"thread": "Mission-type orders in DDIL — what really works?", "depth": "evaluation", "body": "MCDP 6 Ch 3 says mission-type orders. In DDIL it's not enough — you need pre-briefed contingencies because the orient step in OODA stalls when comms degrade. The doctrine is right in spirit but soft on the EW reality."},
        {"thread": "Casualty evac — squad-level decision authority", "depth": "synthesis", "body": "The doctrine pushes CASEVAC launch to platoon. I'd push it to squad. The few minutes saved are decisive. That's a doctrine + policy conversation."},
        {"thread": "Reverse-slope defense — when?", "depth": "analysis", "body": "When the enemy has overhead ISR or armor, reverse slope removes their direct observation. MCWP 3-01 covers this but the modern ISR threat makes it more important than the doctrine implies."},
        {"thread": "Squad fire and movement — the basics", "depth": "application", "body": "Per MCRP 3-10A.4: bounding overwatch with team-level fire support. We rehearsed last week and it worked."},
        {"thread": "Mentorship in MCRP 6-11D — what stuck?", "depth": "recall", "body": "The five horizontal themes — fidelity, fortitude, initiative, perseverance, integrity. I keep them on my notebook cover."},
        {"thread": "Tactical patience — when do you wait vs press?", "depth": "evaluation", "body": "Doctrine doesn't give you a clean answer. My rule: press when you have initiative and intel. Wait when you have neither. Doctrine implies always press; that's wrong against a peer."},
        {"thread": "Decision games — getting the most out of them", "depth": "analysis", "body": "Most TDGs ask 'what would you do?' The harder ones ask 'what would you do at T+15s, T+45s, T+2m?' Doctrine doesn't teach you to think on multiple time horizons."},
        {"thread": "Heat casualty mitigation — leader actions", "depth": "application", "body": "Per the heat cat SOPs: water plan briefed, work-rest cycles enforced, a buddy-check at the top of every hour. We had a near-miss last summer because I missed a buddy check window."},
        {"thread": "How do you brief commander's intent for a junior?", "depth": "synthesis", "body": "Strip it to: end state, two key tasks, one thing you absolutely don't want. That's it. Anything more and the junior won't carry it forward under stress."},
    ],
    "M003": [
        {"thread": "PMCS — what's the difference between scheduled and corrective?", "depth": "recall", "body": "Scheduled is on a calendar. Corrective is when something breaks. From MCO 4790.25."},
        {"thread": "GCSS-MC requisition flow — walk through", "depth": "recall", "body": "You log in, you make the requisition, it routes. I'm still learning the screens."},
        {"thread": "Class IX demand forecast — how do you do it?", "depth": "recall", "body": "Look at past consumption. Project forward. I think there's a tool in GCSS for this."},
        {"thread": "Why is PMCS important?", "depth": "recall", "body": "It catches problems before they become deadlines. Keeps equipment ready."},
        {"thread": "Recovery operations — squad responsibilities", "depth": "recall", "body": "Per MCWP 4-11.4 — we have to recover or evac. Cross-level if we can."},
        {"thread": "What's the RREL?", "depth": "recall", "body": "Readiness reportable equipment list. From MCO 4790.25 Ch 4."},
        {"thread": "Echelons of maintenance — confused on field vs sustainment", "depth": "recall", "body": "Field is at the operating unit. Sustainment is higher. I keep forgetting which is which."},
        {"thread": "Class V planning — never done it before", "depth": "recall", "body": "I haven't done a Class V plan. Asking for help. Anyone done one for a rifle company?"},
        {"thread": "GCSS-MC outage — what do we do?", "depth": "application", "body": "I think we go to manual requisitions. The SOP is in the binder somewhere. I have to find it."},
        {"thread": "Why did my PMCS Quiz score so low?", "depth": "recall", "body": "I missed the questions on calendar-based intervals. Going to re-read MCO 4790.25 Ch 2 this week."},
    ],
}


# ---------------------------------------------------------------------------
# Synthetic assignment instructions, rubrics, and student submissions
# ---------------------------------------------------------------------------

ASSIGNMENT_INSTRUCTIONS = {
    "log_principles_paper": {
        "title": "Logistics Principles Paper — Assignment Instructions",
        "paragraphs": [
            "PURPOSE: Demonstrate the application of the seven principles of logistics (per MCDP 4 Ch 2) to a contemporary expeditionary scenario of your choosing.",
            "TASK: Write a 1,500-word analytical paper. Select a real or notional MAGTF operation in the last 25 years. Identify two of the seven principles that were most stressed. Cite at least three doctrinal sources (MCDP 4, one of MCWP 4-11 / MCWP 4-11.7 / MCRP 4-11.1F, and one of your choosing).",
            "SUBMISSION: Upload a .docx no later than the due date. Use the provided rubric (xlsx) as your self-check before submission.",
            "GRADING: Per the attached rubric.xlsx — five axes scored 0-5, weighted equally.",
            "ACADEMIC INTEGRITY: All work must be your own. Cite all doctrinal references.",
        ],
    },
    "sergeants_course": {
        "title": "Sergeants Course — Squad in the Defense Written Check Instructions",
        "paragraphs": [
            "PURPOSE: Demonstrate Sergeants-Course-level proficiency in defensive planning per MCRP 3-10A.4 Ch 5 and MCWP 3-01 Ch 5. T&R event ties: NAVMC 3500.18 INF-MAN-1001, INF-OPS-2001.",
            "TASK: Given the attached scenario sketch, produce a written defensive plan covering: sectors of fire, primary/alternate/supplementary positions, EAs, withdrawal criteria, and casualty plan.",
            "FORMAT: 800-1,200 words. Reference at least one MCDP and the squad-level doctrine.",
            "GRADING: Per the rubric.xlsx (5 axes, 0-5). Records governance: Privacy Act of 1974 (5 U.S.C. § 552a) and DoDI 1322.35 \"Military Education Records\".",
        ],
    },
    "mos_0411_sustainment": {
        "title": "Class IX Demand Forecast — Practical Exercise",
        "paragraphs": [
            "PURPOSE: Demonstrate the ability to construct a 30-day Class IX (repair parts) demand forecast for a notional rifle battalion per MCWP 4-11.7 Ch 4 and MCO 4790.25.",
            "TASK: Using the historical consumption data provided (xlsx tab 'consumption'), forecast 30-day demand by NSN for the top 50 line items. Identify the top three risk lines (longest expected ASL gap).",
            "FORMAT: Submit a 600-900 word memo + the completed worksheet. Cite MCO 4400.16K UMMIPS priority assignment for any lines pushed to a higher priority designator.",
            "GRADING: Per the attached rubric.xlsx (5 axes, 0-5).",
        ],
    },
}


# Synthetic submissions — one per student matched to their primary course.
# These intentionally reflect the student's profile so the LLM analysis has
# something interesting to grade.
SUBMISSIONS = {
    "M001": {
        "course_id": "log_principles_paper",
        "title": "Logistics Principles Paper — Capt Alvarez, J.",
        "paragraphs": [
            "INTRODUCTION. The seven principles of logistics defined in MCDP 4 chapter 2 — responsiveness, simplicity, flexibility, economy, attainability, sustainability, and survivability — apply across the range of military operations. This paper analyzes the 26th MEU's notional Operation EAGLE THUNDER to argue that responsiveness and survivability were the two most-stressed principles, and that the doctrine's treatment of survivability has not kept pace with the contested-logistics environment.",
            "RESPONSIVENESS DURING EAGLE THUNDER. The CLB pushed 5-day pulses to the BLT in lieu of waiting for pull-based requisition because the supported unit's operational tempo could not absorb a 96-hour requisition lead time. Per MCWP 4-11 Ch 5, the sustainment estimate must be tied to the commander's COA selection. In our case the planner made the COA selection contingent on the CLB's pulse posture; this is a defensible inversion of the doctrinal sequence.",
            "SURVIVABILITY UNDER PERSISTENT ISR. MCDP 4 treats survivability as primarily a hardening question. In a drone-saturated environment, survivability is a distribution question: smaller, more numerous nodes with shorter dwell times. The doctrine implicitly assumes a permissive rear; that assumption fails in 2026. I would argue MCDP 4 needs an addendum on distributed sustainment.",
            "RECOMMENDATION. The Marine Corps should commission a doctrinal update to MCDP 4 Ch 2 explicitly addressing distributed sustainment in contested logistics environments. The update should integrate with MCWP 4-12 Ch 3 on theater distribution.",
            "CONCLUSION. The seven principles remain valid but require modernization. Responsiveness and survivability are the most-stressed in contemporary expeditionary operations.",
        ],
        "self_assessed_areas_of_concern": [
            "Did not cite MCRP 4-11.1F sustainment planning factors directly",
            "Critical thinking section may be thin — only one COA example",
        ],
    },
    "M002": {
        "course_id": "sergeants_course",
        "title": "Sergeants Course (Resident, SNCO Academy) — Squad in the Defense Written Check — Sgt Brennan, K.",
        "paragraphs": [
            "MISSION. 1st Squad, 2nd Platoon, Bravo Co, defends Battle Position 3 NLT 220500ZAPR26 IOT prevent enemy advance along Avenue of Approach NORTH.",
            "CONCEPT. Per MCRP 3-10A.4 Ch 5 and MCWP 3-01 Ch 5: 1st Squad establishes a defense in depth with primary positions oriented NW, alternate positions oriented N to cover EA TIGER, and supplementary positions oriented W to address penetration along Route BLUE.",
            "FIRES. Sectors of fire: 1st Fire Team primary northwest, alternate north. 2nd Fire Team primary north, supplementary west. 3rd Fire Team in reserve, prepared to counter-attack along Route GOLD. Overlap with 2nd Squad at PL RED — 1st Squad owns the seam.",
            "CASUALTY PLAN. CCP at vicinity grid AS 12345 67890. Litter teams pre-staged with 2nd Fire Team. Decision authority for CASEVAC launch held at squad — I am collapsing the OODA loop deliberately because per MCDP 6 Ch 3 mission-type orders require the squad leader to act on commander's intent without a permission cycle.",
            "WITHDRAWAL CRITERIA. We withdraw to alternate positions on commander's order or on 50% combat-effectiveness loss; whichever first. Pre-briefed signal: green star cluster.",
            "DOCTRINAL NOTE. MCDP 1 Ch 4 emphasizes tempo, focus, surprise, boldness, combined arms. In the defense, tempo is denied to the enemy by depth. This plan generates depth via three position arrays and a counter-attack option.",
        ],
        "self_assessed_areas_of_concern": [
            "Writing is terse — not enough explanation for the reader new to the AO",
            "Did not cite a specific MCRP for the casualty plan",
        ],
    },
    "M003": {
        "course_id": "mos_0411_sustainment",
        "title": "Class IX Demand Forecast — PE — Cpl Chen, L.",
        "paragraphs": [
            "PURPOSE. This memo forecasts 30-day Class IX demand for a notional rifle battalion using the historical consumption data provided.",
            "METHOD. I averaged the past 90 days of consumption from the xlsx and projected forward 30 days. I used a flat projection because I was not sure how to do a trend line in the worksheet.",
            "RESULTS. The top 50 NSNs are listed in the worksheet tab DEMAND. The top three risk lines are: (1) NSN 5180-01-XXX-XXXX hex set, (2) NSN 5905-01-XXX-XXXX resistor pack, (3) NSN 6135-01-XXX-XXXX battery 12V.",
            "PRIORITY DESIGNATORS. I left all lines at the routine priority designator. I am not sure when to push something to a higher PD.",
            "WEAKNESSES IN MY OWN WORK. I know I did not cite MCO 4400.16K. I did not really understand the UMMIPS section. I'm going to read MCO 4400.16K Ch 2 this week.",
        ],
        "self_assessed_areas_of_concern": [
            "Did not cite MCO 4400.16K UMMIPS priority designators",
            "Used flat projection — no trend analysis",
            "Did not articulate why these three lines are 'risk' (just longest gap?)",
        ],
    },
}


# ---------------------------------------------------------------------------
# .docx + .xlsx writers (best-effort — degrades gracefully if libs missing)
# ---------------------------------------------------------------------------

def _write_docx(path: Path, title: str, paragraphs: list[str]) -> None:
    """Write a .docx with python-docx. Falls back to .txt if not installed."""
    try:
        from docx import Document
    except Exception:
        path = path.with_suffix(".txt")
        path.write_text(f"{title}\n\n" + "\n\n".join(paragraphs))
        return
    doc = Document()
    doc.add_heading(title, level=1)
    for p in paragraphs:
        doc.add_paragraph(p)
    doc.save(str(path))


def _write_rubric_xlsx(path: Path, course: dict) -> None:
    """Write the rubric as a .xlsx via openpyxl. Falls back to .csv."""
    try:
        from openpyxl import Workbook
    except Exception:
        rows = ["axis,description,weight"]
        for axis in course["rubric_axes"]:
            rows.append(f"{axis},Score 0-5 against published USMC training standards,1.0")
        path.with_suffix(".csv").write_text("\n".join(rows))
        return
    wb = Workbook()
    ws = wb.active
    ws.title = "rubric"
    ws.append(["axis", "0 - Unsatisfactory", "1", "2 - Developing", "3 - Proficient", "4", "5 - Distinguished", "weight"])
    descriptions = {
        "doctrinal_grounding":            ("No doctrinal cites", "1 cite, partial",   "2 cites",          "3+ cites, mostly correct", "3+ cites, accurate", "5+ cites, woven into argument"),
        "logistics_principles_application": ("Principles missing", "Listed not applied", "Applied to one case", "Applied to scenario", "Critically evaluated", "Synthesizes principles + critique"),
        "writing_clarity":                ("Incoherent", "Many errors", "Clear sentences", "Coherent paragraphs", "Polished prose", "Publishable quality"),
        "critical_thinking":              ("Recall only", "Some application", "Application", "Analysis", "Synthesis", "Evaluation"),
        "rubric_alignment":               ("Does not address rubric", "Partially addresses", "Addresses most", "Addresses all", "Addresses + extends", "Exceeds rubric"),
        "leadership_judgment":            ("No judgment exhibited", "Surface judgment", "Defensible decisions", "Sound judgment", "Mature judgment", "Mentor-grade judgment"),
        "tactical_decision_making":       ("Missing DPs", "DPs implied", "DPs stated", "DPs with criteria", "DPs with branches", "DPs with branches + sequels"),
        "maintenance_planning_factors":   ("No factors used", "Factors named", "Factors applied", "Factors applied accurately", "Factors + sensitivity analysis", "Factors + improvement proposal"),
        "systems_thinking":               ("No systems view", "Components named", "Components linked", "Causal links", "Feedback loops identified", "Designs an intervention"),
    }
    for axis in course["rubric_axes"]:
        d = descriptions.get(axis, ("0", "1", "2", "3", "4", "5"))
        ws.append([axis, d[0], d[1], d[2], d[3], d[4], d[5], 1.0])
    wb.save(str(path))


# ---------------------------------------------------------------------------
# Main generators
# ---------------------------------------------------------------------------

def generate(*, seed: int = 1776) -> None:
    """Generate everything under data/."""
    rng = random.Random(seed)  # noqa: F841 (reserved for future randomization)

    # courses.json
    (ROOT / "courses.json").write_text(json.dumps({"courses": COURSES}, indent=2))

    # students.json — pair forum posts to each student
    students_with_posts = []
    for s in STUDENTS:
        sid = s["student_id"]
        students_with_posts.append({
            **s,
            "forum_posts": FORUM_POSTS[sid],
        })
    (ROOT / "students.json").write_text(
        json.dumps({"students": students_with_posts}, indent=2)
    )

    # doctrine_index.json
    (ROOT / "doctrine_index.json").write_text(json.dumps(DOCTRINE_INDEX, indent=2))

    # assignments/*.docx + rubrics/*.xlsx
    for c in COURSES:
        cid = c["id"]
        ai = ASSIGNMENT_INSTRUCTIONS.get(cid, {"title": c["name"], "paragraphs": []})
        _write_docx(ASSIGNMENTS_DIR / f"{cid}_instructions.docx",
                    ai["title"], ai["paragraphs"])
        _write_rubric_xlsx(RUBRICS_DIR / f"{cid}_rubric.xlsx", c)

    # submissions/*.docx (per student)
    for sid, sub in SUBMISSIONS.items():
        _write_docx(SUBMISSIONS_DIR / f"{sid}_submission.docx",
                    sub["title"], sub["paragraphs"])

    # cached_briefs.json — pre-computed per (student_id, course_id) so the
    # demo's hero brief is instant. The text is the deterministic baseline
    # study plan — agent.write_study_plan() will overlay the live LLM brief
    # when the user clicks Regenerate.
    briefs = {}
    for s in STUDENTS:
        sid = s["student_id"]
        cid = s["primary_course_id"]
        course = next(c for c in COURSES if c["id"] == cid)
        analysis = _baseline_analysis(s, course)
        plan = _baseline_study_plan(s, course, analysis)
        briefs[f"{sid}_{cid}"] = {
            "student_id": sid,
            "course_id": cid,
            "analysis": analysis,
            "study_plan": plan,
            "generated_at": datetime.now(timezone.utc).isoformat(),
            "source": "baseline-precompute",
        }
    (ROOT / "cached_briefs.json").write_text(json.dumps(briefs, indent=2))


# ---------------------------------------------------------------------------
# Deterministic baseline analysis + study plan (used as fallback AND seed)
# ---------------------------------------------------------------------------

DEPTH_TO_SCORE = {
    "recall": 1.0, "application": 2.5, "analysis": 3.5,
    "synthesis": 4.5, "evaluation": 5.0,
}


def baseline_analysis(student: dict, course: dict) -> dict:
    """Deterministic per-student adaptive analysis matching the chat_json
    schema. Used as fallback so the UI never sits empty."""
    return _baseline_analysis(student, course)


def _baseline_analysis(student: dict, course: dict) -> dict:
    posts = student.get("forum_posts", [])
    profile = student.get("profile", "developing")
    depths = [p["depth"] for p in posts]
    depth_score = sum(DEPTH_TO_SCORE.get(d, 1.0) for d in depths) / max(1, len(depths))

    # Heuristic gap detection by profile + course
    gaps_by_profile = {
        "developing": ["RSOI throughput calculations", "Sustainment estimate integration with MCPP"],
        "rising_star": ["Written argument structure", "Doctrine citation density"],
        "needs_remediation": ["Class IX planning factors", "UMMIPS priority designator selection",
                              "PMCS calendar-based intervals", "GCSS-MC requisition flow"],
    }
    gaps = gaps_by_profile.get(profile, ["Critical thinking depth"])

    # Doctrinal references "to review" — pull from the course's primary doctrine
    review = course["primary_doctrine"][:2]
    if profile == "needs_remediation":
        review = course["primary_doctrine"]

    # Cite count — count distinct cite-shaped substrings in posts (rough proxy)
    import re as _re
    cite_re = _re.compile(r"MC[A-Z]{2,4}\s*\d+(?:[-\.]\d+)*[A-Z]?(?:\s*Ch\s*\d+)?", _re.I)
    correct_cites = sum(len(cite_re.findall(p["body"])) for p in posts)

    # Writing competency — mean of post body length & profile bias
    avg_words = sum(len(p["body"].split()) for p in posts) / max(1, len(posts))
    writing_base = {
        "developing": 3.0, "rising_star": 4.0, "needs_remediation": 2.0,
    }.get(profile, 3.0)
    writing_score = round(min(5, max(0, writing_base + (avg_words - 30) / 50.0)), 1)

    # Critical-thinking indicators
    indicator_pool = {
        "developing":        ["weighs trade-offs", "applies doctrine to context", "recognizes doctrinal tension"],
        "rising_star":       ["synthesizes across publications", "challenges doctrinal assumptions",
                              "thinks across time horizons"],
        "needs_remediation": ["recall-level engagement", "names doctrine without applying it"],
    }
    indicators = indicator_pool.get(profile, ["application-level engagement"])

    # Recommended study questions — tailored to gaps
    questions_by_profile = {
        "developing": [
            "Walk through an RSOI throughput calculation for a notional MEU APOD with two C-17 spots and a 22-hour ops day. Show your work.",
            "Where does the sustainment estimate enter the MCPP cycle, and what happens if the planner skips it?",
            "Pick one of MCDP 4's seven principles and argue why it is most stressed in distributed maritime operations.",
        ],
        "rising_star": [
            "Rewrite your defensive plan with one additional explicit decision point covering the seam with 2nd Squad.",
            "Compare MCDP 6 mission-type orders against your own DDIL-environment experience. Where does doctrine fall short?",
            "Brief your commander's intent for an upcoming TDG in three sentences. End with the one thing you absolutely don't want to happen.",
        ],
        "needs_remediation": [
            "Define the difference between scheduled and corrective maintenance per MCO 4790.25 Ch 2. Cite the chapter.",
            "Walk through a Class IX requisition in GCSS-MC step by step.",
            "When would you push a UMMIPS line from priority designator 09 to 03? Cite MCO 4400.16K Ch 2.",
        ],
    }
    questions = questions_by_profile.get(profile, [])

    # Peer learning suggestions — pick the highest-depth posts in the cohort
    peer_threads = [p["thread"] for p in posts if p["depth"] in ("synthesis", "evaluation", "analysis")][:3]
    if not peer_threads:
        peer_threads = [p["thread"] for p in posts[:3]]

    # Estimated competency alignment percentage
    base_pct = {"developing": 70, "rising_star": 85, "needs_remediation": 55}.get(profile, 70)

    return {
        "student_id": student["student_id"],
        "course_id": course["id"],
        "knowledge_gaps_identified": gaps,
        "doctrinal_references_cited_correctly": int(correct_cites),
        "doctrinal_references_to_review": review,
        "writing_competency_score": writing_score,
        "critical_thinking_indicators": indicators,
        "recommended_study_questions": questions,
        "peer_learning_suggestions": peer_threads,
        "estimated_competency_alignment_pct": base_pct,
        "cognitive_depth_observed": (
            "evaluation" if depth_score >= 4.5 else
            "synthesis"  if depth_score >= 3.8 else
            "analysis"   if depth_score >= 3.0 else
            "application" if depth_score >= 2.0 else "recall"
        ),
        "_source": "baseline",
    }


def baseline_study_plan(student: dict, course: dict, analysis: dict) -> str:
    """Public alias used by agent.py."""
    return _baseline_study_plan(student, course, analysis)


def _baseline_study_plan(student: dict, course: dict, analysis: dict) -> str:
    name = student["name"]
    course_name = course["name"]
    gaps = analysis["knowledge_gaps_identified"]
    review = analysis["doctrinal_references_to_review"]
    questions = analysis["recommended_study_questions"]
    peer = analysis["peer_learning_suggestions"]
    pct = analysis["estimated_competency_alignment_pct"]
    cd = analysis["cognitive_depth_observed"]

    days = []
    # Seed the 7-day plan
    day_targets = [
        ("Day 1 — Doctrine re-grounding", review, "Read the cited chapters; write a 100-word abstract for each in your own words."),
        ("Day 2 — Gap drill", gaps, "Pick the top gap. Write a 1-page brief addressing it. Cite at least two doctrinal sources."),
        ("Day 3 — Practice question", [questions[0]] if questions else [], "Spend 45 minutes drafting a written response. Self-grade against the rubric."),
        ("Day 4 — Peer learning", peer[:2], "Read these forum threads. Reply to at least one with a synthesis-level post."),
        ("Day 5 — Practice question", [questions[1]] if len(questions) > 1 else [], "Repeat day-3 protocol. Compare your two drafts."),
        ("Day 6 — Rubric self-check", course["rubric_axes"], "Score your draft against each rubric axis 0-5. Identify one axis to improve before submission."),
        ("Day 7 — Capstone", [questions[2]] if len(questions) > 2 else gaps[:1], "Submit your final draft to the instructor. Note three doctrinal references in your conclusion."),
    ]
    for header, items, action in day_targets:
        days.append(f"### {header}")
        if items:
            for it in items:
                days.append(f"- {it}")
        days.append(f"*Action:* {action}\n")

    rubric_tips = "\n".join(
        f"- **{axis.replace('_',' ').title()}**: address this rubric axis explicitly in the body of your draft."
        for axis in course["rubric_axes"]
    )

    tr_manual = course.get("tr_manual", "")
    tr_codes = ", ".join(course.get("tr_event_codes", []))
    tr_line = (
        f"**T&R / PME authority:** {tr_manual}"
        + (f"  ·  **T&R event codes:** {tr_codes}" if tr_codes else "")
        + "  \n"
    ) if tr_manual else ""

    return (
        f"# Adaptive Study Plan — {name}\n"
        f"**Course:** {course_name} ({course['code']})  \n"
        f"{tr_line}"
        f"**Estimated competency alignment:** {pct}%  \n"
        f"**Cognitive depth observed:** {cd}\n\n"
        f"## Knowledge Gaps Identified\n"
        + "\n".join(f"- {g}" for g in gaps) + "\n\n"
        f"## Doctrinal References to Review\n"
        + "\n".join(f"- **{r}** — {DOCTRINE_INDEX.get(r, {}).get('title', '(see catalog)')}" for r in review)
        + "\n\n"
        f"## 7-Day Learning Targets\n"
        + "\n".join(days)
        + f"\n## Rubric-Aligned Writing Tips\n{rubric_tips}\n\n"
        f"---\n"
        f"_Originator: CADENCE — Adaptive PME Tutor. "
        f"Classification: **UNCLASSIFIED // FOR OFFICIAL USE — Military Education Records**.  \n"
        f"Records governance: **Privacy Act of 1974 (5 U.S.C. § 552a) and DoDI 1322.35 "
        f"\"Military Education Records\"** (NOT FERPA — FERPA does not apply to "
        f"active-duty military training).  \n"
        f"Audit chain: SHA-256 chained append-only log under data/audit_logs/cadence_audit.jsonl. "
        f"This Marine's submissions, forum posts, and the analysis above never leave the accredited environment._\n"
    )


def main() -> None:
    generate()
    print("CADENCE data generated:")
    for f in sorted(ROOT.rglob("*")):
        if f.is_file():
            print(f"  {f.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
